import pdfplumber
import docx
import csv
import os
from io import StringIO, BytesIO
from PIL import Image
import base64


def extract_text_tables(file_path: str) -> str:
    result = ""
    if file_path.endswith(".pdf"):
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    result += "<p>" + text.replace("\n", "</p><p>") + "</p>"

                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        csv_output = StringIO()
                        csv_writer = csv.writer(csv_output)
                        csv_writer.writerows(table)
                        result += f"<pre>{csv_output.getvalue()}</pre>"

                # Извлечение изображений
                if page.images:
                    for img in page.images:
                        img_data = img["stream"].get_data()
                        encoded_img = base64.b64encode(img_data).decode("utf-8")
                        result += f'<img src="data:image/png;base64,{encoded_img}"/>'

    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)

        text_data = []
        table_data = []
        image_data = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_data.append(f"<p>{para.text}</p>")

        for table in doc.tables:
            csv_output = StringIO()
            csv_writer = csv.writer(csv_output)
            for row in table.rows:
                csv_writer.writerow([cell.text.strip() for cell in row.cells])
            table_data.append(f"<pre>{csv_output.getvalue()}</pre>")

        # Извлечение изображений
        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                image_data_blob = doc.part.rels[rel].target_part.blob
                encoded_img = base64.b64encode(image_data_blob).decode("utf-8")
                image_data.append(f'<img src="data:image/png;base64,{encoded_img}"/>')

        if text_data:
            result += "".join(text_data)
        if table_data:
            result += "".join(table_data)
        if image_data:
            result += "".join(image_data)

    return result
